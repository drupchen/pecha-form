"""Style designer (Phase 4). Booklet typography is data-driven: ORGANIZATION-wide role
templates + per-DOCUMENT overrides, resolved default ← org ← document by the frontend at
render time (the built-in defaults live in the frontend so bench + print stay identical).
Organizations/users are not built yet — everything is org 1 for now.
"""
import io
import json
from typing import Optional

from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from fastapi.responses import Response
from pydantic import BaseModel

from ..db import get_db

router = APIRouter(prefix="/api", tags=["styles"])

DEFAULT_ORG = 1

# The documented style-name CONVENTION — role → (canonical docx paragraph-style name,
# accepted aliases incl. the reference booklet's names). Used by both the template
# export and the import so a template round-trips transparently.
ROLE_STYLE_NAMES: dict[str, tuple[str, list[str]]] = {
    "tibetan_body":  ("Tibetan Body", ["བོད་ཡིག", "Tibetan"]),
    # Tibetan integrated on the recto (above the phonetics), not left on the verso.
    "tibetan_inline": ("Tibetan In Translation",
                       ["བོད་ཡིག in Translation", "ཡིག་ཆེན།", "ཡིག་ཆེན། in Translation"]),
    "tibetan_title": ("Tibetan Section Title", ["ས་བཅད།", "Tibetan Section"]),
    "tibetan_small": ("Tibetan Small Letters", ["ཡིག་ཆུང་།", "ཡིག་ཆུང"]),
    "phonetics":     ("Phonetics", ["Phonetics Words", "Phonetics With-Tib"]),
    "translation":   ("Translation", ["Translation Words", "Translation With-Tib"]),
    "mantra":        ("Mantras", ["Mantra", "Mantras Words"]),
    "small":         ("Small Letters", ["Small Words", "Small"]),
    "intro":         ("Introduction", ["Intro"]),
    # Section titles by outline level (three tiers). Legacy single "Sections"/"Section"
    # imports onto level 1.
    "section_1":     ("Section Title 1", ["Sections", "Section", "Section Title"]),
    "section_2":     ("Section Title 2", []),
    "section_3":     ("Section Title 3", []),
    "title_tib":     ("Title Tibetan", ["ཁ་བྱང་།", "ཁ་བྱང"]),
    "title_main":    ("Title", ["Text Title", "Book Title", "Main Title"]),
    "title_sub":     ("Subtitle", ["Sub Title"]),
    "copyright":     ("Copyright", []),
    "toc":           ("Table of Contents", ["TOC", "Contents"]),
    "folio":         ("Folio", ["Page Number", "Footer"]),
    "image_caption": ("Caption", ["Image Caption"]),
}

# Concrete role defaults for a starter template — GROUNDED in the reference docx's named
# styles (font/size/bold/italic/indent pt), used when neither org nor document has a value.
_TEMPLATE_DEFAULTS: dict[str, dict] = {
    "tibetan_body":  {"font": "Chogyal", "size": 16},
    "tibetan_inline": {"font": "Chogyal", "size": 16},                          # བོད་ཡིག in Translation
    "tibetan_title": {"font": "Chogyal", "size": 11, "indent": 11.3},           # ས་བཅད
    "tibetan_small": {"font": "Chogyal", "size": 12},                            # ཡིག་ཆུང
    "phonetics":     {"font": "Raleway SemiBold", "size": 10, "indent": 28.4},   # Phonetics
    "translation":   {"font": "Gentium Basic", "size": 11, "indent": 42.5},      # Translation
    "mantra":        {"font": "Gentium Basic", "size": 12, "bold": True, "indent": 28.4},  # Mantras (Words)
    "small":         {"font": "Libertinus Serif Display", "size": 9},            # Small Letters
    "intro":         {"font": "Gentium Basic", "size": 11},                      # Introduction

    "section_1":     {"font": "Libertinus Serif Display", "size": 15},           # Sections L1 (upright)
    "section_2":     {"font": "Libertinus Serif Display", "size": 13.5},         # Sections L2
    "section_3":     {"font": "Libertinus Serif Display", "size": 12},           # Sections L3
    "title_tib":     {"font": "Chogyal", "size": 24},                            # ཁ་བྱང
    "title_main":    {"font": "Libertinus Serif Semibold", "size": 18},          # Title
    "title_sub":     {"font": "Calibri", "size": 12, "italic": True},            # Subtitle
    "copyright":     {"font": "Gentium Basic", "size": 11},
    "toc":           {"font": "Gentium Basic", "size": 11},
    "folio":         {"font": "Georgia", "size": 9},
    "image_caption": {"font": "Gentium Basic", "size": 12, "italic": True},      # Caption
}
ALLOWED_FONT_MIME = {"font/ttf", "font/otf", "font/woff", "font/woff2",
                     "application/font-sfnt", "application/x-font-ttf",
                     "application/x-font-otf", "application/octet-stream"}
MAX_FONT_BYTES = 8 * 1024 * 1024


class StyleProps(BaseModel):
    props: dict


class OrgLayoutIn(BaseModel):
    # Page format / guides, in mm. Partial: only the keys being changed.
    config: dict


class StyleRow(BaseModel):
    role: str
    props: dict


class FontOut(BaseModel):
    id: int
    family: str
    weight: int
    italic: bool
    mime: str


def _rows_to_map(rows) -> dict:
    out = {}
    for r in rows:
        try:
            out[r["role"]] = json.loads(r["props"] or "{}")
        except json.JSONDecodeError:
            out[r["role"]] = {}
    return out


# ── Organization-wide role styles ────────────────────────────────────────────

@router.get("/styles")
def get_org_styles(org_id: int = DEFAULT_ORG):
    conn = get_db()
    try:
        rows = conn.execute(
            "SELECT role, props FROM style_roles WHERE org_id = ?", (org_id,)).fetchall()
        return _rows_to_map(rows)
    finally:
        conn.close()


@router.put("/styles/{role}")
def put_org_style(role: str, payload: StyleProps, org_id: int = DEFAULT_ORG):
    conn = get_db()
    try:
        conn.execute(
            "INSERT INTO style_roles (org_id, role, props) VALUES (?, ?, ?) "
            "ON CONFLICT(org_id, role) DO UPDATE SET props = excluded.props",
            (org_id, role, json.dumps(payload.props)))
        conn.commit()
        return {"role": role, "props": payload.props}
    finally:
        conn.close()


@router.delete("/styles/{role}")
def delete_org_style(role: str, org_id: int = DEFAULT_ORG):
    conn = get_db()
    try:
        conn.execute("DELETE FROM style_roles WHERE org_id = ? AND role = ?", (org_id, role))
        conn.commit()
        return {"ok": True}
    finally:
        conn.close()


# ── Organization-wide page format and guides ─────────────────────────────────
# The sheet size and the four margins the text block and the binding/folio guides are drawn
# from. Resolved default ← org ← document by `documents._effective_config`, so a booklet that
# says nothing about its geometry follows the house and one that does keeps its own.

@router.get("/org-layout")
def get_org_layout(org_id: int = DEFAULT_ORG):
    """The org's page format. Always complete: the built-in geometry, then whatever the org
    has said, so the caller never has to render the word "inherit"."""
    from .documents import DEFAULT_LAYOUT_CONFIG, ORG_LAYOUT_KEYS
    conn = get_db()
    try:
        cfg = {k: DEFAULT_LAYOUT_CONFIG[k] for k in ORG_LAYOUT_KEYS}
        row = conn.execute("SELECT config FROM org_layout WHERE org_id = ?", (org_id,)).fetchone()
        if row and row["config"]:
            try:
                stored = json.loads(row["config"])
                if isinstance(stored, dict):
                    cfg.update({k: v for k, v in stored.items() if k in ORG_LAYOUT_KEYS})
            except (ValueError, TypeError):
                pass
        return cfg
    finally:
        conn.close()


@router.put("/org-layout")
def put_org_layout(payload: OrgLayoutIn, org_id: int = DEFAULT_ORG):
    """Merge geometry onto the org template. Unknown keys are dropped rather than stored: the
    template speaks for the page format and the guides, not for type sizes (the roles own
    those) or for the bench's reflow delay."""
    from .documents import DEFAULT_LAYOUT_CONFIG, ORG_LAYOUT_KEYS
    conn = get_db()
    try:
        row = conn.execute("SELECT config FROM org_layout WHERE org_id = ?", (org_id,)).fetchone()
        cfg = {k: DEFAULT_LAYOUT_CONFIG[k] for k in ORG_LAYOUT_KEYS}
        if row and row["config"]:
            try:
                cfg.update(json.loads(row["config"]))
            except (ValueError, TypeError):
                pass
        for k, v in payload.config.items():
            if k not in ORG_LAYOUT_KEYS:
                continue
            try:
                f = float(v)
            except (TypeError, ValueError):
                raise HTTPException(400, f"{k} must be a number")
            if not (f > 0):
                raise HTTPException(400, f"{k} must be greater than 0")
            cfg[k] = f
        cfg = {k: cfg[k] for k in ORG_LAYOUT_KEYS}   # complete, and in a stable order
        conn.execute(
            "INSERT INTO org_layout (org_id, config) VALUES (?, ?) "
            "ON CONFLICT(org_id) DO UPDATE SET config = excluded.config",
            (org_id, json.dumps(cfg)))
        conn.commit()
        return cfg
    finally:
        conn.close()


# ── Per-document overrides ───────────────────────────────────────────────────

@router.get("/documents/{document_id}/styles")
def get_doc_styles(document_id: int):
    conn = get_db()
    try:
        rows = conn.execute(
            "SELECT role, props FROM document_style_overrides WHERE document_id = ?",
            (document_id,)).fetchall()
        return _rows_to_map(rows)
    finally:
        conn.close()


@router.put("/documents/{document_id}/styles/{role}")
def put_doc_style(document_id: int, role: str, payload: StyleProps):
    conn = get_db()
    try:
        if not conn.execute("SELECT 1 FROM documents WHERE id = ?", (document_id,)).fetchone():
            raise HTTPException(404, "Document not found")
        conn.execute(
            "INSERT INTO document_style_overrides (document_id, role, props) VALUES (?, ?, ?) "
            "ON CONFLICT(document_id, role) DO UPDATE SET props = excluded.props",
            (document_id, role, json.dumps(payload.props)))
        conn.commit()
        return {"role": role, "props": payload.props}
    finally:
        conn.close()


@router.delete("/documents/{document_id}/styles/{role}")
def delete_doc_style(document_id: int, role: str):
    conn = get_db()
    try:
        conn.execute(
            "DELETE FROM document_style_overrides WHERE document_id = ? AND role = ?",
            (document_id, role))
        conn.commit()
        return {"ok": True}
    finally:
        conn.close()


# ── Organization fonts (@font-face) ──────────────────────────────────────────

@router.get("/org-fonts")
def list_org_fonts(org_id: int = DEFAULT_ORG):
    conn = get_db()
    try:
        rows = conn.execute(
            "SELECT id, family, weight, italic, mime FROM org_fonts WHERE org_id = ? "
            "ORDER BY family, weight, italic", (org_id,)).fetchall()
        return [FontOut(id=r["id"], family=r["family"], weight=r["weight"],
                        italic=bool(r["italic"]), mime=r["mime"]) for r in rows]
    finally:
        conn.close()


@router.post("/org-fonts")
async def upload_org_font(
    file: UploadFile = File(...),
    family: str = Form(...),
    weight: int = Form(400),
    italic: bool = Form(False),
    org_id: int = Form(DEFAULT_ORG),
):
    data = await file.read()
    mime = file.content_type or "application/octet-stream"
    if mime not in ALLOWED_FONT_MIME:
        raise HTTPException(415, f"Unsupported font type: {mime}")
    if not data:
        raise HTTPException(400, "Empty file")
    if len(data) > MAX_FONT_BYTES:
        raise HTTPException(413, "Font too large (max 8 MB)")
    if not family.strip():
        raise HTTPException(400, "family is required")
    conn = get_db()
    try:
        cur = conn.execute(
            "INSERT INTO org_fonts (org_id, family, weight, italic, mime, data) "
            "VALUES (?, ?, ?, ?, ?, ?)",
            (org_id, family.strip(), weight, int(bool(italic)), mime, data))
        conn.commit()
        return FontOut(id=cur.lastrowid, family=family.strip(), weight=weight,
                       italic=bool(italic), mime=mime)
    finally:
        conn.close()


@router.get("/org-fonts/{font_id}/file")
def get_org_font_file(font_id: int):
    conn = get_db()
    try:
        row = conn.execute("SELECT mime, data FROM org_fonts WHERE id = ?", (font_id,)).fetchone()
        if not row:
            raise HTTPException(404, "Font not found")
        return Response(content=row["data"], media_type=row["mime"],
                        headers={"Cache-Control": "public, max-age=86400"})
    finally:
        conn.close()


@router.delete("/org-fonts/{font_id}")
def delete_org_font(font_id: int):
    conn = get_db()
    try:
        conn.execute("DELETE FROM org_fonts WHERE id = ?", (font_id,))
        conn.commit()
        return {"ok": True}
    finally:
        conn.close()


# ── Organization cover seal (the ༀ placeholder's image) ──────────────────────
# Part of the TEMPLATE, like the fonts: it prints on every booklet's cover, where the ༀ
# ornament otherwise sits. A booklet that uploads its own cover image overrides it.

ALLOWED_IMAGE_MIME = {"image/png", "image/jpeg", "image/webp", "image/gif"}
MAX_IMAGE_BYTES = 10 * 1024 * 1024


class SealOut(BaseModel):
    has_image: bool = False
    width_mm: Optional[float] = None
    height_mm: Optional[float] = None


class SealSizeIn(BaseModel):
    width_mm: Optional[float] = None    # NULL/absent = the image's natural size
    height_mm: Optional[float] = None


@router.get("/org-seal", response_model=SealOut)
def get_org_seal(org_id: int = DEFAULT_ORG):
    conn = get_db()
    try:
        row = conn.execute(
            "SELECT width_mm, height_mm FROM org_seal WHERE org_id = ?", (org_id,)).fetchone()
        if not row:
            return SealOut()
        return SealOut(has_image=True, width_mm=row["width_mm"], height_mm=row["height_mm"])
    finally:
        conn.close()


@router.get("/org-seal/file")
def get_org_seal_file(org_id: int = DEFAULT_ORG):
    conn = get_db()
    try:
        row = conn.execute(
            "SELECT mime, data FROM org_seal WHERE org_id = ?", (org_id,)).fetchone()
        if not row:
            raise HTTPException(404, "No seal for this organization")
        return Response(content=row["data"], media_type=row["mime"],
                        headers={"Cache-Control": "no-store"})
    finally:
        conn.close()


@router.put("/org-seal", response_model=SealOut)
async def upload_org_seal(file: UploadFile = File(...), org_id: int = Form(DEFAULT_ORG)):
    data = await file.read()
    mime = file.content_type or ""
    if mime not in ALLOWED_IMAGE_MIME:
        raise HTTPException(415, f"Unsupported image type: {mime or 'unknown'}")
    if not data:
        raise HTTPException(400, "Empty file")
    if len(data) > MAX_IMAGE_BYTES:
        raise HTTPException(413, "Image too large (max 10 MB)")
    conn = get_db()
    try:
        # A replacement keeps the size the designer already set (only the bytes change).
        conn.execute(
            "INSERT INTO org_seal (org_id, mime, data) VALUES (?, ?, ?) "
            "ON CONFLICT(org_id) DO UPDATE SET mime = excluded.mime, data = excluded.data, "
            "updated_at = CURRENT_TIMESTAMP",
            (org_id, mime, data))
        conn.commit()
        row = conn.execute(
            "SELECT width_mm, height_mm FROM org_seal WHERE org_id = ?", (org_id,)).fetchone()
        return SealOut(has_image=True, width_mm=row["width_mm"], height_mm=row["height_mm"])
    finally:
        conn.close()


@router.patch("/org-seal", response_model=SealOut)
def set_org_seal_size(payload: SealSizeIn, org_id: int = DEFAULT_ORG):
    conn = get_db()
    try:
        row = conn.execute("SELECT 1 FROM org_seal WHERE org_id = ?", (org_id,)).fetchone()
        if not row:
            raise HTTPException(404, "No seal for this organization")
        conn.execute(
            "UPDATE org_seal SET width_mm = ?, height_mm = ?, updated_at = CURRENT_TIMESTAMP "
            "WHERE org_id = ?", (payload.width_mm, payload.height_mm, org_id))
        conn.commit()
        return SealOut(has_image=True, width_mm=payload.width_mm, height_mm=payload.height_mm)
    finally:
        conn.close()


@router.delete("/org-seal")
def delete_org_seal(org_id: int = DEFAULT_ORG):
    conn = get_db()
    try:
        conn.execute("DELETE FROM org_seal WHERE org_id = ?", (org_id,))
        conn.commit()
        return {"ok": True}
    finally:
        conn.close()


# ── Style Studio specimen (per-org editable sample) ──────────────────────────

class SampleIn(BaseModel):
    content: str


@router.get("/style-sample")
def get_style_sample(org_id: int = DEFAULT_ORG):
    conn = get_db()
    try:
        row = conn.execute(
            "SELECT content FROM style_samples WHERE org_id = ?", (org_id,)).fetchone()
        return {"content": row["content"] if row else ""}
    finally:
        conn.close()


@router.put("/style-sample")
def put_style_sample(payload: SampleIn, org_id: int = DEFAULT_ORG):
    conn = get_db()
    try:
        conn.execute(
            "INSERT INTO style_samples (org_id, content) VALUES (?, ?) "
            "ON CONFLICT(org_id) DO UPDATE SET content = excluded.content",
            (org_id, payload.content))
        conn.commit()
        return {"ok": True}
    finally:
        conn.close()


# ── docx style templates (import / export) ───────────────────────────────────
# An external docx whose NAMED styles follow ROLE_STYLE_NAMES seeds org or document
# styles — parsed to derive font family/size/weight/italic/colour/alignment. The
# starter template names its styles by the convention so parsing is transparent.

def _props_from_docx_style(style) -> dict:
    """Map a python-docx paragraph style → our StyleProps (only what is set)."""
    props: dict = {}
    f = style.font
    if f.name:
        props["fontFamily"] = f.name
    if f.size is not None:
        props["fontSize"] = f"{f.size.pt:g}pt"
    if f.bold:                       # only explicit bold sets weight (semibold lives in the family name)
        props["fontWeight"] = 700
    if f.italic is not None:
        props["italic"] = bool(f.italic)
    # COLOUR IS NOT IMPORTED. The booklet prints black by default (see bookletStyles.ts);
    # a docx's screen colours (Word's blue heading tints) would otherwise arrive as org-wide
    # defaults nobody chose. Any other ink is set by hand in the Style Studio.
    # Paragraph styles carry alignment/spacing/indent; character styles don't.
    pf = getattr(style, "paragraph_format", None)
    if pf is not None:
        align = pf.alignment
        if align is not None:
            mapped = {0: "left", 1: "center", 2: "right", 3: "justify"}.get(int(align))
            if mapped:
                props["align"] = mapped
        if pf.line_spacing is not None and isinstance(pf.line_spacing, float):
            props["lineHeight"] = f"{pf.line_spacing:g}"
        if pf.left_indent is not None:
            props["indent"] = f"{pf.left_indent.mm:.1f}mm"
    return props


@router.get("/style-template.docx")
def export_style_template(target: str = "document", org_id: int = DEFAULT_ORG,
                          document_id: int | None = None):
    """Generate a starter .docx with one named paragraph style per role, valued from the
    resolved current styles (defaults ← org ← [document])."""
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.style import WD_STYLE_TYPE

    conn = get_db()
    try:
        org = _rows_to_map(conn.execute(
            "SELECT role, props FROM style_roles WHERE org_id = ?", (org_id,)).fetchall())
        doc_ov = {}
        if target == "document" and document_id is not None:
            doc_ov = _rows_to_map(conn.execute(
                "SELECT role, props FROM document_style_overrides WHERE document_id = ?",
                (document_id,)).fetchall())
    finally:
        conn.close()

    out = docx.Document()
    out.add_paragraph("Booklet style template — edit each style's font/size/weight/italic, "
                      "then re-import. Do not rename the styles.").italic = True
    for role, (name, _aliases) in ROLE_STYLE_NAMES.items():
        base = dict(_TEMPLATE_DEFAULTS.get(role, {}))
        # resolve: template default ← org ← document (only keys present override)
        for src in (org.get(role, {}), doc_ov.get(role, {})):
            if src.get("fontFamily"): base["font"] = src["fontFamily"]
            if src.get("fontSize"):
                try: base["size"] = float(str(src["fontSize"]).rstrip("pt"))
                except ValueError: pass
            if "fontWeight" in src: base["bold"] = int(src.get("fontWeight") or 0) >= 600
            if "italic" in src: base["italic"] = bool(src["italic"])
            if src.get("indent"):
                try: base["indent"] = float(str(src["indent"]).rstrip("mm")) / 25.4 * 72
                except ValueError: pass
        try:
            st = out.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            st = out.styles[name]
        if base.get("font"): st.font.name = base["font"]
        if base.get("size"): st.font.size = Pt(float(base["size"]))
        if base.get("bold") is not None: st.font.bold = bool(base.get("bold"))
        if base.get("italic") is not None: st.font.italic = bool(base.get("italic"))
        if base.get("indent"): st.paragraph_format.left_indent = Pt(float(base["indent"]))
        p = out.add_paragraph(f"{name} — sample text  ༄༅།", style=name)  # noqa: RUF001

    buf = io.BytesIO()
    out.save(buf)
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="booklet-style-template.docx"'})


@router.post("/style-template/import")
async def import_style_template(
    file: UploadFile = File(...),
    target: str = Form("document"),
    org_id: int = Form(DEFAULT_ORG),
    document_id: int | None = Form(None),
):
    """Parse an uploaded docx's named styles → role StyleProps → write to org or document."""
    import docx

    if target not in ("org", "document"):
        raise HTTPException(400, "target must be 'org' or 'document'")
    if target == "document" and document_id is None:
        raise HTTPException(400, "document_id is required for target=document")

    data = await file.read()
    try:
        d = docx.Document(io.BytesIO(data))
    except Exception:
        raise HTTPException(400, "Could not read the file as a .docx")

    # name (lowercased) → role, over canonical names + aliases
    name_to_role: dict[str, str] = {}
    for role, (canon, aliases) in ROLE_STYLE_NAMES.items():
        for nm in (canon, *aliases):
            name_to_role[nm.strip().lower()] = role

    # Merge every style matching a role: a CHARACTER style ("… Words") wins on the run's
    # font attrs; the PARAGRAPH style supplies indent/align/line-height. So we get both.
    FONT_KEYS = ("fontFamily", "fontSize", "fontWeight", "italic")   # never colour — see above
    PARA_KEYS = ("align", "lineHeight", "indent")
    applied: dict[str, dict] = {}
    for style in d.styles:
        try:
            role = name_to_role.get((style.name or "").strip().lower())
        except Exception:
            role = None
        if not role:
            continue
        props = _props_from_docx_style(style)
        if not props:
            continue
        cur = applied.setdefault(role, {})
        is_char = getattr(style, "paragraph_format", None) is None
        for k in FONT_KEYS:
            if k in props and (is_char or k not in cur):
                cur[k] = props[k]
        for k in PARA_KEYS:
            if k in props:
                cur[k] = props[k]

    if not applied:
        raise HTTPException(422, "No matching named styles found. See the expected style names.")

    conn = get_db()
    try:
        for role, props in applied.items():
            blob = json.dumps(props)
            if target == "org":
                conn.execute(
                    "INSERT INTO style_roles (org_id, role, props) VALUES (?, ?, ?) "
                    "ON CONFLICT(org_id, role) DO UPDATE SET props = excluded.props",
                    (org_id, role, blob))
            else:
                conn.execute(
                    "INSERT INTO document_style_overrides (document_id, role, props) VALUES (?, ?, ?) "
                    "ON CONFLICT(document_id, role) DO UPDATE SET props = excluded.props",
                    (document_id, role, blob))
        conn.commit()
        return {"applied": list(applied.keys()), "count": len(applied)}
    finally:
        conn.close()
