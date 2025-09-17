from pathlib import Path

from docx import Document # package name: python-docx
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH


class FormatDocumentUpdated:
    def __init__(self, template=None):
        if not template:
            self.template = Path(__file__).parent.resolve() / 'templates' / 'template_Monlam.docx'
        else:
            self.template = Path(template)

        # 1. parse input files
        self.document = self.__initiate_document()

        # styles for the translation
        #     par styles
        self.s_title = 'Text Title'
        self.s_sections = 'Sections'
        self.s_trans = 'Translation'
        self.s_sml = 'Small Letters'
        self.s_phon = 'Phonetics'
        self.s_mtr = 'Mantras'
        #    char styles
        self.s_char_trans = 'Translation Words'
        self.s_char_sml = 'Small Words'
        self.s_char_phon = 'Phonetics Words'
        self.s_char_mtr = 'Mantras Words'
        # unused
        #self.s_trans_tib = 'Translation With-Tib'
        #self.s_phon_tib = 'Phonetics With-Tib'
        #self.s_bo_par_trans = 'བོད་ཡིག in Translation'

        # style equivalence tables
        self.bklt_par_styles = {
            'T': self.s_title,
            'sub': self.s_trans,
            't1': self.s_sections,
            't2': self.s_sections,
            'n': self.s_trans,
            's': self.s_sml,
            'k': self.s_mtr
            # b for bold and i for italics are directly in the code
        }
        self.bklt_char_styles = {
            'T': None,
            'sub': None,
            't1': None,  #self.s_char_sml,
            't2': None,  #self.s_char_sml,
            'n': self.s_char_trans,
            's': self.s_char_sml,
            'k': self.s_char_mtr
        }

        # styles for the Tibetan
        #     par styles
        self.s_bo_par = 'བོད་ཡིག'
        self.s_bo_par_title1 = 'ཁ་བྱང་།'
        self.s_bo_par_title2 = 'ས་བཅད།'
        #     char styles
        self.s_bo_char_title1 = 'ཁ་བྱང་ཡི་གེ'
        self.s_bo_char_big = 'ཡིག་ཆེན།'
        self.s_bo_char_small = 'ཡིག་ཆུང་།'
        # style equivalence tables
        self.bo_par_styles = {
            'T': self.s_bo_par_title1,
            't': self.s_bo_par_title2,
            's': self.s_bo_par,
            'n': self.s_bo_par,
            'b': self.s_bo_par,
        }
        self.bo_char_styles = {
            'T': self.s_bo_char_title1,
            't': self.s_bo_char_small,
            's': self.s_bo_char_small,
            'n': self.s_bo_char_big,
            'b': self.s_bo_char_big,
        }

    def format_booklet(self, parsed_content, out_file, no_phon=False):
        def is_long(string):
            if len(string.split(' ')) >= 15:
                return True
            return False

        for s, parts in parsed_content:
            for part in parts:
                # if there is phonetic or sanskrit
                if part['tib'] or part['skt']:
                    if part['skt']:
                        par = self.document.add_paragraph(style=self.s_mtr)
                        run = par.add_run(part['skt'], style=self.s_char_mtr)
                        run.bold = True
                        if is_long(part['skt']): par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    if not no_phon and part['phon']:
                        par = self.document.add_paragraph(style=self.s_phon)
                        par.add_run(part['phon'], style=self.s_char_phon)
                        if is_long(part['phon']): par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    # hack: remove empty translations. should be done in the parsing function
                    if len(part['trans']) == 1 and not part['trans'][0]:
                        part['trans'].pop()

                # translation
                if not part['trans']:  # pass empty translations
                    continue
                if not part['trans_formatted'][0][0] and s == 'k':
                    continue

                par = self.document.add_paragraph(style=self.bklt_par_styles[s])
                #if s == 's':
                #    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for string in part['trans_formatted']:
                    if isinstance(string, tuple):
                        string, style = string
                        run = par.add_run(string)
                        for stl, boolean in style.items():
                            if stl in self.bklt_char_styles:
                                run.style = self.bklt_char_styles[stl]
                            elif boolean:
                                if stl == 'bold':
                                    run.bold = True
                                if stl == 'italic':
                                    run.italic = True
                    else:
                        par.add_run(string, style=self.bklt_char_styles[s])

                if is_long(par.text): par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # add page break after Title
            if s == 'T':
                par.runs[-1].add_break(WD_BREAK.PAGE)

        self.document.save(out_file)

    def format_tibetan(self, parsed_content, out_file):
        for t, strings in parsed_content:
            par = self.document.add_paragraph(style=self.bo_par_styles[t])
            for s in strings:
                text, style = s
                if style['small letters']:
                    p = 's'
                else:
                    p = t
                par.add_run(text, style=self.bo_char_styles[p])
            if t == 'T':
                par.runs[-1].add_break(WD_BREAK.PAGE)
        self.document.save(out_file)

    def __initiate_document(self):
        def delete_paragraph(paragraph):
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None

        doc = Document(self.template)
        for par in doc.paragraphs:
            delete_paragraph(par)
        return doc
